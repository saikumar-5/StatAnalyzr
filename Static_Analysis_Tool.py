import hashlib
import requests
import os
import time
import subprocess
import json
import re
from pathlib import Path
import zipfile
import shutil
import platform

# --- IMPORTANT: PLACE YOUR VIRUSTOTAL API KEY HERE ---
# Replace 'YOUR_VIRUSTOTAL_API_KEY_HERE' with your actual VirusTotal API Key.
VIRUSTOTAL_API_KEY = '6fa8bd53cf29d967f03114ee6ff3140a2d16de8769b596dc3f465043f7546ef3'
# -----------------------------------------------------

try:
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
except ImportError:
    print("python-docx not found. Installing...")
    subprocess.run(['pip', 'install', 'python-docx'], check=True)
    import docx
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE

class StaticMalwareAnalyzer:
    def __init__(self):
        self.virustotal_api_key = VIRUSTOTAL_API_KEY
        self.tools_dir = Path('./tools')
        self.floss_path = self.tools_dir / 'floss'
        self.upx_path = self.tools_dir / 'upx'
        self.setup_tools()

    def setup_tools(self):
        print("Setting up analysis tools...")
        self.tools_dir.mkdir(parents=True, exist_ok=True)

        is_windows = platform.system() == "Windows"

        if not self.floss_path.exists():
            print("Downloading FLOSS...")
            floss_url = "https://github.com/mandiant/flare-floss/releases/download/v3.0.0/floss-v3.0.0-linux.zip"
            floss_zip_path = self.tools_dir / 'floss.zip'
            try:
                with requests.get(floss_url, stream=True) as r:
                    r.raise_for_status()
                    with open(floss_zip_path, 'wb') as f:
                        for chunk in r.iter_content(chunk_size=8192):
                            f.write(chunk)
                with zipfile.ZipFile(floss_zip_path, 'r') as zip_ref:
                    zip_ref.extractall(self.tools_dir)
                if not is_windows:
                    os.chmod(self.floss_path, 0o755)
                print("FLOSS installed successfully")
            except Exception as e:
                print(f"Error downloading or extracting FLOSS: {e}")
            finally:
                if floss_zip_path.exists():
                    os.remove(floss_zip_path)

        if not self.upx_path.exists():
            print("Downloading UPX...")
            upx_url = 'https://github.com/upx/upx/releases/download/v4.0.2/upx-4.0.2-amd64_linux.tar.xz'
            upx_tar_path = self.tools_dir / 'upx.tar.xz'
            try:
                with requests.get(upx_url, stream=True) as r:
                    r.raise_for_status()
                    with open(upx_tar_path, 'wb') as f:
                        for chunk in r.iter_content(chunk_size=8192):
                            f.write(chunk)
                
                import tarfile
                with tarfile.open(upx_tar_path, 'r:xz') as tar_ref:
                    for member in tar_ref.getmembers():
                        if member.name.endswith('/upx') and member.isfile():
                            with open(self.upx_path, 'wb') as outfile, tar_ref.extractfile(member) as infile:
                                outfile.write(infile.read())
                            break
                if not is_windows:
                    os.chmod(self.upx_path, 0o755)
                print("UPX installed successfully")
            except Exception as e:
                print(f"Error downloading or extracting UPX: {e}")
            finally:
                if upx_tar_path.exists():
                    os.remove(upx_tar_path)

        try:
            try:
                import pefile
            except ImportError:
                print("pefile not found. Installing...")
                subprocess.run(['pip', 'install', 'pefile'], check=True, capture_output=True)

            try:
                import yara
            except ImportError:
                print("yara-python not found. Installing...")
                subprocess.run(['pip', 'install', 'yara-python'], check=True, capture_output=True)

            print("PE analysis libraries ensured.")
        except Exception as e:
            print(f"Error installing PE libraries: {e}")

    def sanitize_string(self, s):
        if not isinstance(s, str):
            s = str(s)

        sanitized = ''.join(char for char in s if ord(char) >= 32 or char in '\t\n\r')

        if len(sanitized) > 1000:
            sanitized = sanitized[:1000] + "... [truncated]"

        return sanitized

    def calculate_hashes(self, filepath):
        md5_hash = hashlib.md5()
        sha1_hash = hashlib.sha1()
        sha256_hash = hashlib.sha256()

        try:
            with open(filepath, "rb") as f:
                for byte_block in iter(lambda: f.read(4096), b""):
                    md5_hash.update(byte_block)
                    sha1_hash.update(byte_block)
                    sha256_hash.update(byte_block)
        except FileNotFoundError:
            print(f"Error: File not found at {filepath}")
            return None, None, None
        except Exception as e:
            print(f"Error calculating hashes for {filepath}: {e}")
            return None, None, None

        return md5_hash.hexdigest(), sha1_hash.hexdigest(), sha256_hash.hexdigest()

    def query_virustotal(self, file_hash, hash_type):
        if not self.virustotal_api_key or self.virustotal_api_key == 'YOUR_VIRUSTOTAL_API_KEY_HERE':
            print("VirusTotal API key is not configured. Skipping VirusTotal query.")
            return None

        url = f"https://www.virustotal.com/api/v3/files/{file_hash}"
        headers = {"x-apikey": self.virustotal_api_key}

        print(f"Querying VirusTotal for {hash_type} hash: {file_hash}")

        try:
            response = requests.get(url, headers=headers, timeout=30)

            if response.status_code == 200:
                return response.json()
            elif response.status_code == 404:
                return {"response_code": 0, "message": "File not found in VirusTotal database"}
            else:
                return {"error": f"API returned status code {response.status_code}: {response.text}"}
        except requests.exceptions.RequestException as e:
            return {"error": f"Request failed: {str(e)}"}

    def analyze_with_floss(self, filepath):
        print("\n" + "="*50)
        print("STEP 2: FLOSS STRING EXTRACTION")
        print("="*50)

        if not self.floss_path.exists():
            print("FLOSS executable not found. Skipping string extraction.")
            return None

        try:
            cmd = [str(self.floss_path), '--json', str(filepath)]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=300, errors='ignore')

            if result.returncode == 0:
                try:
                    floss_data = json.loads(result.stdout)
                    self.display_floss_results(floss_data)
                    return floss_data
                except json.JSONDecodeError:
                    print("FLOSS output (raw):")
                    print(result.stdout[:2000] + "..." if len(result.stdout) > 2000 else result.stdout)
                    return {"raw_output": result.stdout}
            else:
                print(f"FLOSS failed with return code {result.returncode}")
                print(f"Error: {result.stderr}")
                return None

        except subprocess.TimeoutExpired:
            print("FLOSS analysis timed out. This might happen with large or complex binaries.")
            return None
        except Exception as e:
            print(f"Error running FLOSS: {e}")
            return None

    def display_floss_results(self, floss_data):
        if 'strings' in floss_data:
            strings = floss_data['strings']

            print(f"Found {len(strings)} strings:")

            urls = []
            files = []
            registry = []
            apis = []
            other = []

            for string_item in strings:
                if isinstance(string_item, dict):
                    s = string_item.get('string', '')
                else:
                    s = str(string_item)

                if re.match(r'https?://', s):
                    urls.append(s)
                elif '\\' in s and ('.' in s or 'HKEY' in s.upper()):
                    if 'HKEY' in s.upper():
                        registry.append(s)
                    else:
                        files.append(s)
                elif (s.endswith('.dll') or s.endswith('.exe')) and len(s) > 4:
                    apis.append(s)
                elif len(s) > 4 and s.isprintable():
                    other.append(s)

            if urls:
                print(f"\nURLs found ({len(urls)}):")
                for url in urls[:10]:
                    print(f"  - {url}")

            if files:
                print(f"\nFile paths found ({len(files)}):")
                for file in files[:10]:
                    print(f"  - {file}")

            if registry:
                print(f"\nRegistry keys found ({len(registry)}):")
                for reg in registry[:10]:
                    print(f"  - {reg}")

            if apis:
                print(f"\nPotential APIs/DLLs ({len(apis)}):")
                for api in apis[:10]:
                    print(f"  - {api}")

            if other:
                print(f"\nOther interesting strings ({len(other)}):")
                for s in other[:20]:
                    if len(s) > 5:
                        print(f"  - {s}")

    def analyze_with_upx(self, filepath):
        print("\n" + "="*50)
        print("STEP 3: UPX PACKING ANALYSIS")
        print("="*50)

        if not self.upx_path.exists():
            print("UPX executable not found. Skipping packing analysis.")
            return None, {"status": "UPX tool not found."}

        try:
            info_cmd = [str(self.upx_path), '-t', str(filepath)]
            info_result = subprocess.run(info_cmd, capture_output=True, text=True, errors='ignore')

            if "IsPacked: Yes" in info_result.stdout or info_result.returncode == 0:
                print("File appears to be packed with UPX")
                print("UPX Info:")
                print(info_result.stdout)
                packing_status = "Packed with UPX"
                upx_info_output = info_result.stdout

                unpacked_path = filepath.parent / (filepath.name + "_unpacked")
                unpack_cmd = [str(self.upx_path), '-d', str(filepath), '-o', str(unpacked_path)]
                unpack_result = subprocess.run(unpack_cmd, capture_output=True, text=True, errors='ignore')

                if unpack_result.returncode == 0 and unpacked_path.exists():
                    print(f"Successfully unpacked to: {unpacked_path}")

                    original_size = os.path.getsize(filepath)
                    unpacked_size = os.path.getsize(unpacked_path)

                    print(f"Original size: {original_size} bytes")
                    print(f"Unpacked size: {unpacked_size} bytes")
                    print(f"Compression ratio: {original_size/unpacked_size:.2f}x")

                    return str(unpacked_path), {
                        "status": "Unpacked",
                        "original_size": original_size,
                        "unpacked_size": unpacked_size,
                        "compression_ratio": f"{original_size/unpacked_size:.2f}x",
                        "upx_info": upx_info_output,
                        "unpacked_filepath": str(unpacked_path)
                    }
                else:
                    print("Failed to unpack file")
                    print(unpack_result.stderr)
                    return None, {
                        "status": "Packed, but failed to unpack",
                        "upx_info": upx_info_output,
                        "unpack_error": unpack_result.stderr
                    }
            else:
                print("File is not packed with UPX or UPX cannot handle this file")
                return None, {"status": "Not UPX packed or not detectable", "upx_info": info_result.stdout}

        except Exception as e:
            print(f"Error running UPX analysis: {e}")
            return None, {"status": "Error during UPX analysis", "error": str(e)}

    def analyze_pe_structure(self, filepath):
        print("\n" + "="*50)
        print("STEP 4: PE STRUCTURE ANALYSIS")
        print("="*50)

        pe_analysis_data = {}

        try:
            import pefile

            pe = pefile.PE(str(filepath))

            pe_analysis_data['basic_info'] = {
                'machine_type': hex(pe.FILE_HEADER.Machine),
                'number_of_sections': pe.FILE_HEADER.NumberOfSections,
                'time_date_stamp': time.ctime(pe.FILE_HEADER.TimeDateStamp),
                'entry_point': hex(pe.OPTIONAL_HEADER.AddressOfEntryPoint),
                'image_base': hex(pe.OPTIONAL_HEADER.ImageBase)
            }
            print("=== BASIC PE INFORMATION ===")
            for k, v in pe_analysis_data['basic_info'].items():
                print(f"{k.replace('_', ' ').title()}: {v}")

            pe_analysis_data['sections'] = []
            print("\n=== SECTION ANALYSIS ===")
            for section in pe.sections:
                section_info = {
                    'name': section.Name.decode(errors='ignore').strip(),
                    'virtual_address': hex(section.VirtualAddress),
                    'virtual_size': section.Misc_VirtualSize,
                    'raw_size': section.SizeOfRawData,
                    'characteristics': hex(section.Characteristics),
                    'warnings': []
                }
                print(f"Section: {section_info['name']}")
                print(f"  Virtual Address: {section_info['virtual_address']}")
                print(f"  Virtual Size: {section_info['virtual_size']}")
                print(f"  Raw Size: {section_info['raw_size']}")
                print(f"  Characteristics: {section_info['characteristics']}")

                if section.SizeOfRawData == 0:
                    section_info['warnings'].append("Section has zero raw size (often indicates obfuscation)")
                    print("  ⚠️ WARNING: Section has zero raw size")
                if section.Misc_VirtualSize > section.SizeOfRawData * 2 and section.SizeOfRawData != 0:
                    section_info['warnings'].append("Virtual size much larger than raw size (possible packing/unpacking anomaly)")
                    print("  ⚠️ WARNING: Virtual size much larger than raw size")

                pe_analysis_data['sections'].append(section_info)

            pe_analysis_data['imports'] = []
            print("\n=== IMPORT ANALYSIS ===")
            if hasattr(pe, 'DIRECTORY_ENTRY_IMPORT'):
                suspicious_apis = [
                    'CreateProcess', 'WriteProcessMemory', 'VirtualAlloc', 'CreateRemoteThread',
                    'LoadLibrary', 'GetProcAddress', 'RegSetValue', 'CreateFile', 'InternetOpen',
                    'HttpSendRequest', 'CreateService', 'SetWindowsHook', 'NtMapViewOfSection',
                    'AdjustTokenPrivileges', 'LookupPrivilegeValue', 'OpenProcessToken',
                    'ShellExecute', 'URLDownloadToFile', 'WinExec', 'CreateRemoteThread'
                ]

                found_suspicious_apis = []
                imported_libraries = set()

                for entry in pe.DIRECTORY_ENTRY_IMPORT:
                    dll_name = entry.dll.decode(errors='ignore')
                    imported_libraries.add(dll_name)
                    dll_imports = []
                    print(f"\nDLL: {dll_name}")

                    for imp in entry.imports:
                        if imp.name:
                            api_name = imp.name.decode(errors='ignore')
                            dll_imports.append(api_name)
                            print(f"  - {api_name}")

                            if any(sus_api.lower() in api_name.lower() for sus_api in suspicious_apis):
                                found_suspicious_apis.append(f"{dll_name}!{api_name}")
                    pe_analysis_data['imports'].append({'dll': dll_name, 'functions': dll_imports})

                if found_suspicious_apis:
                    print("\n⚠️ SUSPICIOUS API CALLS DETECTED:")
                    for api in found_suspicious_apis:
                        print(f"  - {api}")
                pe_analysis_data['suspicious_apis'] = found_suspicious_apis
                pe_analysis_data['imported_libraries'] = sorted(list(imported_libraries))

            pe_analysis_data['exports'] = []
            if hasattr(pe, 'DIRECTORY_ENTRY_EXPORT'):
                print("\n=== EXPORT ANALYSIS ===")
                export_dll_name = pe.DIRECTORY_ENTRY_EXPORT.name.decode(errors='ignore')
                num_exports = len(pe.DIRECTORY_ENTRY_EXPORT.symbols)
                print(f"Export DLL Name: {export_dll_name}")
                print(f"Number of Exports: {num_exports}")
                pe_analysis_data['export_info'] = {'dll_name': export_dll_name, 'num_exports': num_exports}

                exports_list = []
                for exp in pe.DIRECTORY_ENTRY_EXPORT.symbols:
                    if exp.name:
                        exports_list.append(exp.name.decode(errors='ignore'))
                        print(f"  - {exp.name.decode(errors='ignore')}")
                pe_analysis_data['exports'] = exports_list

            print("\n=== PACKER DETECTION ===")
            packer_detection_results = self.detect_packers(pe)
            pe_analysis_data['packer_detection'] = packer_detection_results

            pe.close()
            return pe_analysis_data

        except ImportError:
            print("pefile library not available. Please ensure it's installed and rerun.")
            return {"error": "pefile library not found."}
        except Exception as e:
            print(f"Error analyzing PE structure: {e}")
            return {"error": str(e)}

    def detect_packers(self, pe):
        detected_packers = []
        packer_info = {}

        section_names_bytes = [section.Name.strip(b'\x00') for section in pe.sections]
        section_names_str = [name.decode(errors='ignore') for name in section_names_bytes]

        if any(name.lower() in [b'upx0', b'upx1', b'upx2'] for name in section_names_bytes):
            detected_packers.append('UPX')
        if b'.aspack' in section_names_bytes:
            detected_packers.append('ASPack')
        if b'.petite' in section_names_bytes:
            detected_packers.append('Petite')
        if b'.mpress1' in section_names_bytes or b'.mpress2' in section_names_bytes:
            detected_packers.append('MPRESS')
        if b'Themida' in section_names_bytes:
            detected_packers.append('Themida')
        if b'VMProtect' in section_names_bytes:
            detected_packers.append('VMProtect')
        if b'Fsg' in section_names_bytes:
            detected_packers.append('FSG')

        packer_info['detected_by_sections'] = detected_packers if detected_packers else "None"

        if detected_packers:
            print("Detected packers/obfuscators (by section names):", ', '.join(detected_packers))
        else:
            print("No common packer signatures found in section names.")

        high_entropy_sections = []
        for section in pe.sections:
            data = section.get_data()
            if len(data) > 0:
                entropy = self.calculate_entropy(data)
                if entropy > 7.0:
                    high_entropy_sections.append(f"{section.Name.decode(errors='ignore').strip()} (Entropy: {entropy:.2f})")

        packer_info['high_entropy_sections'] = high_entropy_sections

        if high_entropy_sections:
            print(f"High entropy sections (possible packing/encryption): {', '.join(high_entropy_sections)}")
        else:
            print("No unusually high entropy sections detected.")

        return packer_info

    def calculate_entropy(self, data):
        if len(data) == 0:
            return 0

        frequency = {}
        for byte in data:
            frequency[byte] = frequency.get(byte, 0) + 1

        entropy = 0
        length = len(data)
        import math
        for count in frequency.values():
            probability = count / length
            entropy -= probability * math.log2(probability)

        return entropy

    def generate_report(self, filepath, analysis_results):
        print("\n" + "="*50)
        print("GENERATING COMPREHENSIVE ANALYSIS REPORT (Word Document)")
        print("="*50)

        document = docx.Document()
        document.add_heading('Malware Static Analysis Report', level=1)

        document.add_heading('1. File Information', level=2)
        document.add_paragraph(f"Filename: {self.sanitize_string(analysis_results['file_info']['filename'])}")
        document.add_paragraph(f"File Size: {analysis_results['file_info']['size']} bytes")
        document.add_paragraph(f"Analysis Date: {self.sanitize_string(analysis_results['file_info']['analysis_date'])}")

        document.add_heading('2. File Hashes', level=2)
        hash_table = document.add_table(rows=1, cols=2)
        hash_table.rows[0].cells[0].text = 'Hash Type'
        hash_table.rows[0].cells[1].text = 'Hash Value'
        hash_table.style = 'Table Grid'

        for hash_type, hash_value in analysis_results['hashes'].items():
            row_cells = hash_table.add_row().cells
            row_cells[0].text = hash_type.upper()
            row_cells[1].text = self.sanitize_string(hash_value)

        document.add_heading('3. VirusTotal Scan Results', level=2)
        vt_result = analysis_results.get('virustotal', {})
        if vt_result and 'data' in vt_result:
            data = vt_result['data']
            attributes = data.get('attributes', {})
            stats = attributes.get('last_analysis_stats', {})
            document.add_paragraph(f"Last Analysis Date: {time.ctime(attributes.get('last_analysis_date', 0))}")
            document.add_paragraph(f"Number of Detections: {stats.get('malicious', 0)} / {sum(stats.values())}")
            
            # Add VirusTotal Link to the report
            sha256_hash = analysis_results['hashes'].get('sha256')
            if sha256_hash:
                vt_link = f"https://www.virustotal.com/gui/file/{sha256_hash}/detection"
                document.add_paragraph(f"VirusTotal Link: {self.sanitize_string(vt_link)}")

            vt_table = document.add_table(rows=1, cols=2)
            vt_table.rows[0].cells[0].text = 'Category'
            vt_table.rows[0].cells[1].text = 'Count'
            vt_table.style = 'Table Grid'

            for category, count in stats.items():
                row_cells = vt_table.add_row().cells
                row_cells[0].text = category.replace('_', ' ').title()
                row_cells[1].text = str(count)
        else:
            document.add_paragraph("VirusTotal results not available or file not found in database.")
            if vt_result and 'error' in vt_result:
                document.add_paragraph(f"Error: {self.sanitize_string(str(vt_result['error']))}")

        document.add_heading('4. FLOSS String Extraction', level=2)
        floss_data = analysis_results.get('floss', {})
        if floss_data:
            if 'strings' in floss_data:
                strings = floss_data['strings']
                document.add_paragraph(f"Total deobfuscated strings found: {len(strings)}")
                urls = []
                files = []
                registry = []
                apis = []
                other = []

                for string_item in strings:
                    if isinstance(string_item, dict):
                        s = string_item.get('string', '')
                    else:
                        s = str(string_item)

                    if re.match(r'https?://', s):
                        urls.append(s)
                    elif '\\' in s and ('.' in s or 'HKEY' in s.upper()):
                        if 'HKEY' in s.upper():
                            registry.append(s)
                        else:
                            files.append(s)
                    elif (s.endswith('.dll') or s.endswith('.exe')) and len(s) > 4:
                        apis.append(s)
                    elif len(s) > 4 and s.isprintable():
                        other.append(s)

                if urls:
                    document.add_heading('4.1. URLs Found', level=3)
                    for url in urls:
                        document.add_paragraph(f"- {self.sanitize_string(url)}")
                if files:
                    document.add_heading('4.2. File Paths Found', level=3)
                    for file_path in files:
                        document.add_paragraph(f"- {self.sanitize_string(file_path)}")
                if registry:
                    document.add_heading('4.3. Registry Keys Found', level=3)
                    for reg_key in registry:
                        document.add_paragraph(f"- {self.sanitize_string(reg_key)}")
                if apis:
                    document.add_heading('4.4. Potential APIs/DLLs', level=3)
                    for api in apis:
                        document.add_paragraph(f"- {self.sanitize_string(api)}")
                if other:
                    document.add_heading('4.5. Other Interesting Strings', level=3)
                    for s in other:
                        if len(s) > 5:
                            document.add_paragraph(f"- {self.sanitize_string(s)}")
            elif 'raw_output' in floss_data:
                document.add_paragraph("FLOSS could not parse JSON, showing raw output (first 2000 chars):")
                document.add_paragraph(self.sanitize_string(floss_data['raw_output'][:2000] + "..."))
            else:
                document.add_paragraph("FLOSS data available but no 'strings' or 'raw_output' key found.")
        else:
            document.add_paragraph("FLOSS analysis skipped or failed to produce any results.")

        document.add_heading('5. UPX Packing Analysis', level=2)
        upx_info = analysis_results.get('upx', {})
        if upx_info:
            document.add_paragraph(f"Packing Status: {self.sanitize_string(str(upx_info.get('status', 'N/A')))}")
            if upx_info.get('status') == "Unpacked":
                document.add_paragraph(f"Original Size: {upx_info.get('original_size')} bytes")
                document.add_paragraph(f"Unpacked Size: {upx_info.get('unpacked_size')} bytes")
                document.add_paragraph(f"Compression Ratio: {self.sanitize_string(str(upx_info.get('compression_ratio')))}")
                document.add_paragraph(f"Unpacked File Path: {self.sanitize_string(str(upx_info.get('unpacked_filepath', 'N/A')))}")
                document.add_paragraph("UPX Info Output:")
                document.add_paragraph(self.sanitize_string(str(upx_info.get('upx_info', ''))))
            elif upx_info.get('status') == "Packed, but failed to unpack":
                document.add_paragraph("UPX Info Output:")
                document.add_paragraph(self.sanitize_string(str(upx_info.get('upx_info', ''))))
                document.add_paragraph("Unpack Error:")
                document.add_paragraph(self.sanitize_string(str(upx_info.get('unpack_error', ''))))
            elif upx_info.get('status') == "Not UPX packed or not detectable":
                document.add_paragraph("UPX Info Output:")
                document.add_paragraph(self.sanitize_string(str(upx_info.get('upx_info', ''))))
            else:
                document.add_paragraph(f"Details: {self.sanitize_string(str(upx_info.get('error', 'N/A')))}")
        else:
            document.add_paragraph("UPX analysis skipped or failed.")

        document.add_heading('6. PE Structure Analysis', level=2)
        pe_data = analysis_results.get('pe', {})
        if pe_data and not pe_data.get('error'):
            document.add_heading('6.1. Basic PE Information', level=3)
            for k, v in pe_data.get('basic_info', {}).items():
                document.add_paragraph(f"{k.replace('_', ' ').title()}: {self.sanitize_string(str(v))}")

            document.add_heading('6.2. Section Analysis', level=3)
            section_table = document.add_table(rows=1, cols=6)
            section_table.rows[0].cells[0].text = 'Name'
            section_table.rows[0].cells[1].text = 'Virtual Address'
            section_table.rows[0].cells[2].text = 'Virtual Size'
            section_table.rows[0].cells[3].text = 'Raw Size'
            section_table.rows[0].cells[4].text = 'Characteristics'
            section_table.rows[0].cells[5].text = 'Warnings'
            section_table.style = 'Table Grid'

            for section_info in pe_data.get('sections', []):
                row_cells = section_table.add_row().cells
                row_cells[0].text = self.sanitize_string(section_info.get('name', 'N/A'))
                row_cells[1].text = self.sanitize_string(section_info.get('virtual_address', 'N/A'))
                row_cells[2].text = self.sanitize_string(str(section_info.get('virtual_size', 'N/A')))
                row_cells[3].text = self.sanitize_string(str(section_info.get('raw_size', 'N/A')))
                row_cells[4].text = self.sanitize_string(section_info.get('characteristics', 'N/A'))
                warnings_text = "\n".join([self.sanitize_string(w) for w in section_info.get('warnings', [])])
                row_cells[5].text = warnings_text if warnings_text else "None"

            document.add_heading('6.3. Import Analysis', level=3)
            if pe_data.get('imports'):
                for dll_entry in pe_data['imports']:
                    document.add_paragraph(f"DLL: {self.sanitize_string(dll_entry['dll'])}")
                    for func in dll_entry['functions']:
                        document.add_paragraph(f"  - {self.sanitize_string(func)}")
                if pe_data.get('suspicious_apis'):
                    document.add_paragraph("Suspicious API Calls Detected:", style='Intense Quote')
                    for api in pe_data['suspicious_apis']:
                        document.add_paragraph(f"  - {self.sanitize_string(api)}")
            else:
                document.add_paragraph("No imports found or import directory not present.")

            document.add_heading('6.4. Export Analysis', level=3)
            if pe_data.get('exports'):
                document.add_paragraph(f"Export DLL Name: {self.sanitize_string(pe_data['export_info']['dll_name'])}")
                document.add_paragraph(f"Number of Exports: {pe_data['export_info']['num_exports']}")
                for exp in pe_data['exports']:
                    document.add_paragraph(f"  - {self.sanitize_string(exp)}")
            else:
                document.add_paragraph("No exports found or export directory not present.")

            document.add_heading('6.5. Packer Detection (PE Sections)', level=3)
            packer_detection = pe_data.get('packer_detection', {})
            document.add_paragraph(f"Detected by Section Names: {self.sanitize_string(', '.join(packer_detection.get('detected_by_sections', [])) if packer_detection.get('detected_by_sections') != 'None' else 'None')}")
            if packer_detection.get('high_entropy_sections'):
                document.add_paragraph("High Entropy Sections (possible packing/encryption):")
                for sec in packer_detection['high_entropy_sections']:
                    document.add_paragraph(f"- {self.sanitize_string(sec)}")
            else:
                document.add_paragraph("No unusually high entropy sections detected.")
            
            document.add_heading('6.6. Imported Libraries', level=3)
            imported_libs = pe_data.get('imported_libraries')
            if imported_libs:
                for lib in imported_libs:
                    document.add_paragraph(f"- {self.sanitize_string(lib)}")
            else:
                document.add_paragraph("No imported libraries found.")

        else:
            document.add_paragraph("PE Structure analysis skipped or failed.")
            if pe_data and 'error' in pe_data:
                document.add_paragraph(f"Error: {self.sanitize_string(str(pe_data['error']))}")

        report_filename = f"Malware_Analysis_Report_{Path(filepath).stem}.docx"
        document.save(report_filename)
        print(f"Report saved as: {report_filename}")
        return report_filename

    def analyze_file(self, filepath_str):
        filepath = Path(filepath_str)
        
        if not filepath.exists():
            print(f"Error: The file '{filepath_str}' does not exist. Please provide a valid path.")
            return None, None

        print(f"Starting static malware analysis for: {filepath.name}")
        analysis_results = {
            'file_info': {
                'filename': filepath.name,
                'size': os.path.getsize(filepath),
                'analysis_date': time.strftime("%Y-%m-%d %H:%M:%S")
            }
        }

        print("\n" + "="*50)
        print("STEP 1: HASH CALCULATION AND VIRUSTOTAL QUERY")
        print("="*50)
        md5, sha1, sha256 = self.calculate_hashes(filepath)
        if not all([md5, sha1, sha256]):
            print("Hash calculation failed. Aborting analysis.")
            return analysis_results, None

        analysis_results['hashes'] = {'md5': md5, 'sha1': sha1, 'sha256': sha256}
        print(f"MD5: {md5}\nSHA1: {sha1}\nSHA256: {sha256}")

        vt_results = self.query_virustotal(sha256, "SHA256")
        analysis_results['virustotal'] = vt_results
        if vt_results and 'data' in vt_results:
            attributes = vt_results['data'].get('attributes', {})
            stats = attributes.get('last_analysis_stats', {})
            print(f"\nVirusTotal Detections: {stats.get('malicious', 0)} / {sum(stats.values())}")
            print(f"VirusTotal Link: https://www.virustotal.com/gui/file/{sha256}/detection")
        elif vt_results and 'message' in vt_results:
            print(f"\nVirusTotal: {vt_results['message']}")
        elif vt_results and 'error' in vt_results:
            print(f"\nVirusTotal Error: {vt_results['error']}")
        else:
            print("\nVirusTotal query skipped (API key not set or other issue).")


        floss_output = self.analyze_with_floss(filepath)
        analysis_results['floss'] = floss_output

        unpacked_file_path_str, upx_analysis_info = self.analyze_with_upx(filepath)
        analysis_results['upx'] = upx_analysis_info

        pe_analysis_target_file = Path(unpacked_file_path_str) if unpacked_file_path_str else filepath
        pe_analysis_data = self.analyze_pe_structure(pe_analysis_target_file)
        analysis_results['pe'] = pe_analysis_data

        report_path = self.generate_report(filepath, analysis_results)
        print(f"\nAnalysis complete. Report generated at {report_path}")

        if unpacked_file_path_str and Path(unpacked_file_path_str).exists():
            try:
                os.remove(unpacked_file_path_str)
                print(f"Cleaned up unpacked file: {unpacked_file_path_str}")
            except Exception as e:
                print(f"Error cleaning up unpacked file {unpacked_file_path_str}: {e}")

        return analysis_results, report_path

if __name__ == '__main__':
    analyzer = StaticMalwareAnalyzer()

    print("\nEnter the full path to the malware sample you want to analyze (e.g., C:\\malware\\sample.exe or /home/user/malware/sample):")
    filepath_input = input("File path: ").strip()

    if filepath_input:
        analysis_results, report_file = analyzer.analyze_file(filepath_input)

        if report_file:
            print(f"\nYour analysis report '{report_file}' is ready.")
        else:
            print("\nReport generation failed or was skipped due to earlier errors.")
    else:
        print("No file path provided. Exiting.")